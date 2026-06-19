from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

DARK  = RGBColor(0x0F, 0x17, 0x2A)
BLUE  = RGBColor(0x3B, 0x82, 0xF6)
GRAY  = RGBColor(0x64, 0x74, 0x8B)
LGRAY = RGBColor(0x94, 0xA3, 0xB8)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
RED   = RGBColor(0xDC, 0x26, 0x26)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
AMBER = RGBColor(0xD9, 0x77, 0x06)

def rgb(hex6):
    return RGBColor(int(hex6[0:2],16), int(hex6[2:4],16), int(hex6[4:6],16))

def shade_para(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def shade_cell(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{edge}'); b.set(qn('w:val'), 'none'); tblBorders.append(b)
    tblPr.append(tblBorders)

def add_hyperlink(paragraph, url, text, color_hex='3B82F6', bold=False, size=10):
    part   = paragraph.part
    r_id   = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hl     = OxmlElement('w:hyperlink'); hl.set(qn('r:id'), r_id)
    new_r  = OxmlElement('w:r')
    rPr    = OxmlElement('w:rPr')
    c_el   = OxmlElement('w:color'); c_el.set(qn('w:val'), color_hex); rPr.append(c_el)
    u_el   = OxmlElement('w:u');     u_el.set(qn('w:val'), 'single');  rPr.append(u_el)
    sz_el  = OxmlElement('w:sz');    sz_el.set(qn('w:val'), str(int(size*2))); rPr.append(sz_el)
    if bold:
        b_el = OxmlElement('w:b'); rPr.append(b_el)
    new_r.append(rPr)
    t_el = OxmlElement('w:t'); t_el.text = text; new_r.append(t_el)
    hl.append(new_r)
    paragraph._p.append(hl)

def hr(doc):
    pp = doc.add_paragraph()
    pPr = pp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), 'CBD5E1')
    pBdr.append(bot); pPr.append(pBdr)
    pp.paragraph_format.space_before = Pt(2); pp.paragraph_format.space_after = Pt(6)

def gap(doc, pt=10):
    pp = doc.add_paragraph()
    pp.paragraph_format.space_before = Pt(0); pp.paragraph_format.space_after = Pt(pt)

def p(doc, text='', size=11, bold=False, italic=False, color=None,
      align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=8):
    para = doc.add_paragraph(); para.alignment = align
    para.paragraph_format.space_before = Pt(sb)
    para.paragraph_format.space_after  = Pt(sa)
    if text:
        run = para.add_run(text); run.bold = bold; run.italic = italic
        run.font.size = Pt(size)
        if color: run.font.color.rgb = color
    return para

def section_hdr(doc, text):
    p(doc, text, size=13, bold=True, color=DARK, sb=16, sa=4); hr(doc)

def flow_step(doc, num, color_hex, title, desc, is_last=False):
    ft = doc.add_table(rows=1, cols=2); no_borders(ft)
    nc = ft.rows[0].cells[0]; nc.width = Cm(1.6)
    shade_cell(nc, color_hex)
    np_ = nc.paragraphs[0]; np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    np_.paragraph_format.space_before = Pt(10); np_.paragraph_format.space_after = Pt(10)
    nr = np_.add_run(num); nr.bold=True; nr.font.size=Pt(18); nr.font.color.rgb=WHITE
    cc = ft.rows[0].cells[1]; shade_cell(cc, 'F8FAFC')
    cp1 = cc.paragraphs[0]; cp1.paragraph_format.space_before = Pt(8)
    cr1 = cp1.add_run(title); cr1.bold=True; cr1.font.size=Pt(11.5); cr1.font.color.rgb=DARK
    cp2 = cc.add_paragraph(); cp2.paragraph_format.space_after = Pt(8)
    cr2 = cp2.add_run(desc); cr2.font.size=Pt(10); cr2.font.color.rgb=GRAY
    if not is_last:
        ap = doc.add_paragraph(); ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ap.paragraph_format.space_before = Pt(0); ap.paragraph_format.space_after = Pt(0)
        ar = ap.add_run('↓'); ar.font.size=Pt(14); ar.font.color.rgb=LGRAY


# ═══════════════════════════════════════════════
#  TITLE
# ═══════════════════════════════════════════════
p(doc, '🎙️  PODCAST STUDIO BOOKING SYSTEM', size=26, bold=True, color=DARK,
  align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=4)
p(doc, 'Calendar Architecture, Tier-Based Booking & GoHighLevel Setup', size=13,
  italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
p(doc, 'Three Tiers. Three Studios. One Seamless Booking Experience.', size=14,
  bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
p(doc, 'Prepared by Watts Company  ·  Phase 1 Complete  ·  June 2026', size=9,
  color=LGRAY, align=WD_ALIGN_PARAGRAPH.CENTER, sa=14)

hr(doc)

# ── Two highlighted assumption lines ──────────
note1 = doc.add_paragraph(); note1.alignment = WD_ALIGN_PARAGRAPH.CENTER
note1.paragraph_format.space_before = Pt(6); note1.paragraph_format.space_after = Pt(3)
shade_para(note1, 'FEF9C3')
rn1a = note1.add_run('⚠️  '); rn1a.font.size = Pt(11)
rn1b = note1.add_run('All the numbers which are taken are considered as assumptions.')
rn1b.bold=True; rn1b.font.size=Pt(11); rn1b.font.color.rgb=rgb('92400E')

note2 = doc.add_paragraph(); note2.alignment = WD_ALIGN_PARAGRAPH.CENTER
note2.paragraph_format.space_before = Pt(2); note2.paragraph_format.space_after = Pt(14)
shade_para(note2, 'DBEAFE')
rn2a = note2.add_run('📋  '); rn2a.font.size = Pt(11)
rn2b = note2.add_run('Services taken:  Audio Recording,  Video Podcast,  Live Streaming')
rn2b.bold=True; rn2b.font.size=Pt(11); rn2b.font.color.rgb=rgb('1E40AF')

p(doc,
  'This document outlines the complete GoHighLevel booking system built for the podcast studio. '
  'Three membership tiers — Non-Member, Gold, and VIP — each have their own pricing and a dedicated '
  'booking page. Studios are automatically assigned and protected from double-booking across all tiers simultaneously.',
  size=11, color=GRAY, sb=4, sa=20)


# ═══════════════════════════════════════════════
#  SECTION 1 — OVERVIEW
# ═══════════════════════════════════════════════
section_hdr(doc, '🏷️  SYSTEM OVERVIEW')

t = doc.add_table(rows=1, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.CENTER; no_borders(t)
tiers = [
    ('Tier 1','NON-MEMBER','Pay-per-session at standard rates. No monthly commitment required.','F1F5F9','475569'),
    ('Tier 2','GOLD MEMBER','Monthly membership with discounted studio rates and member benefits.','FEF3C7','D97706'),
    ('Tier 3','VIP MEMBER','Premium membership with the best rates and priority studio access.','EDE9FE','7C3AED'),
]
for i,(num,name,desc,bg,tc) in enumerate(tiers):
    c = t.rows[0].cells[i]; shade_cell(c, bg); c.width = Cm(5.5)
    pp = c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.paragraph_format.space_before=Pt(14)
    rr = pp.add_run(num); rr.bold=True; rr.font.size=Pt(9); rr.font.color.rgb=rgb(tc)
    pp2 = c.add_paragraph(); pp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr2 = pp2.add_run(name); rr2.bold=True; rr2.font.size=Pt(13); rr2.font.color.rgb=rgb(tc)
    pp3 = c.add_paragraph(); pp3.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp3.paragraph_format.space_after=Pt(14)
    rr3 = pp3.add_run(desc); rr3.font.size=Pt(9.5); rr3.font.color.rgb=GRAY

gap(doc, 12)

mt = doc.add_table(rows=1, cols=4)
mt.alignment=WD_TABLE_ALIGNMENT.CENTER; no_borders(mt)
metrics = [('9','Service Calendars','3 types × 3 tiers'),('3','Studios (Rooms)','Studio A, B & C'),
           ('2','Equipment Types','Teleprompter + Encoder'),('3','Booking Pages','One per membership tier')]
for i,(num,label,detail) in enumerate(metrics):
    c = mt.rows[0].cells[i]; shade_cell(c,'EFF6FF')
    pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.paragraph_format.space_before=Pt(12)
    rr=pp.add_run(num); rr.bold=True; rr.font.size=Pt(30); rr.font.color.rgb=BLUE
    pp2=c.add_paragraph(); pp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr2=pp2.add_run(label); rr2.bold=True; rr2.font.size=Pt(9.5); rr2.font.color.rgb=DARK
    pp3=c.add_paragraph(); pp3.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp3.paragraph_format.space_after=Pt(12)
    rr3=pp3.add_run(detail); rr3.font.size=Pt(8.5); rr3.font.color.rgb=LGRAY


# ═══════════════════════════════════════════════
#  SECTION 2 — CUSTOMER BOOKING FLOW (updated)
# ═══════════════════════════════════════════════
gap(doc, 16)
section_hdr(doc, '🗂️  CUSTOMER BOOKING FLOW')
p(doc, 'Complete end-to-end journey — from discovery to confirmed studio booking, including membership purchase via GHL API and tier-locked booking experience.',
  size=10.5, color=GRAY, sb=0, sa=14)

flow_step(doc, '01', '3B82F6', 'Podcaster Discovers the Studio',
    'Via website, social media, referral, or paid ads. Lands on the studio\'s online presence and decides to sign up.')

flow_step(doc, '02', '8B5CF6', 'Signs Up on the Portal',
    'Creates an account with name, email, and phone number. A GHL contact record is automatically created '
    'in the background via GHL API — no manual entry required.')

flow_step(doc, '03', '0EA5E9', 'Logs Into Portal Dashboard',
    'After signup the podcaster logs into the portal. At this point they have no membership tag yet — '
    'they see the 3 available membership products to choose from.')

flow_step(doc, '04', 'F59E0B', 'Views & Purchases a Membership Tier',
    'Portal displays 3 GHL products: Non-Member, Gold, and VIP — each with pricing and benefits. '
    'Podcaster selects a tier and pays via Stripe. On successful payment, the GHL API is called '
    'to update the contact record and apply the correct tag: member-non, member-gold, or member-vip.')

flow_step(doc, '05', '7C3AED', 'GHL API Marks the Contact Tag',
    'The GHL API removes any existing tier tag and applies the new one atomically. '
    'This tag is the single source of truth that drives everything downstream — '
    'the correct service menu, pricing, and access level are all determined by this tag. '
    'Upgrade / Downgrade: if a member upgrades from Gold to VIP, the API removes member-gold and applies member-vip. '
    'Downgrade works the same in reverse. The portal reflects the new tier immediately after tag change.')

flow_step(doc, '06', 'EC4899', 'Portal Shows Tier-Locked Service Menu',
    'Portal reads the GHL contact tag via API. Based on the tag, it displays ONLY that tier\'s service menu URL — '
    'the booking page is completely tier-locked. '
    'member-non → Non-Member booking page (standard rates only). '
    'member-gold → Gold booking page (member rates only). '
    'member-vip → VIP booking page (premium rates only). '
    'A member cannot see or access another tier\'s pricing under any circumstance.')

flow_step(doc, '07', '10B981', 'Selects Service & Time Slot',
    'The tier-locked booking page shows exactly 3 services at their tier\'s pricing: '
    'Audio Recording, Video Podcast, and Live Streaming. '
    'Member picks the service and then picks an available date and time slot.')

flow_step(doc, '08', 'F97316', 'GHL Checks Availability — Automatically',
    'Three simultaneous checks before any slot is displayed: '
    '(1) Is a concurrent booking slot (staff) free? '
    '(2) Is a physical studio room — Studio A, B, or C — available? '
    '(3) Is required equipment available (Teleprompter for Video; Streaming Encoder for Live)? '
    'ALL three must pass. If any one fails, that time slot is hidden from the customer.')

flow_step(doc, '09', '22C55E', 'Booking Confirmed',
    'A studio room is automatically assigned by GHL. Confirmation email + SMS sent instantly to the podcaster. '
    'The booking appears in the client\'s GHL Appointment dashboard with the assigned studio, tier, and service clearly listed.',
    is_last=True)

# Upgrade / Downgrade note box
gap(doc, 12)
ut = doc.add_table(rows=1, cols=1); no_borders(ut)
uc = ut.rows[0].cells[0]; shade_cell(uc, 'F0FDF4')
up1 = uc.paragraphs[0]; up1.paragraph_format.space_before = Pt(12)
ur1 = up1.add_run('🔄  Upgrade / Downgrade Flow')
ur1.bold=True; ur1.font.size=Pt(11.5); ur1.font.color.rgb=GREEN
for line in [
    'Member visits portal → selects a new membership tier (upgrade or downgrade)',
    'Stripe processes the new subscription payment',
    'GHL API removes the old tag (e.g. member-gold) and applies the new tag (e.g. member-vip)',
    'Portal re-reads the updated tag → immediately displays the new tier\'s service menu',
    'Booking page updates in real time — no manual intervention by the studio owner required',
]:
    up2 = uc.add_paragraph()
    up2.paragraph_format.space_after = Pt(3)
    ur2 = up2.add_run(f'  →  {line}')
    ur2.font.size=Pt(10); ur2.font.color.rgb=GRAY
uc.paragraphs[-1].paragraph_format.space_after = Pt(12)


# ═══════════════════════════════════════════════
#  SECTION 3 — CALENDAR TABLE
# ═══════════════════════════════════════════════
gap(doc, 16)
section_hdr(doc, '📅  CALENDAR ARCHITECTURE — 9 SERVICE CALENDARS')
p(doc, '9 service calendars — one for every combination of service type and membership tier. '
       'Each has its own pricing and connects exclusively to its tier\'s booking page.',
  size=10.5, color=GRAY, sb=0, sa=12)

cal_t = doc.add_table(rows=10, cols=5); no_borders(cal_t)
hdrs = ['Calendar ID','Service Type','Tier','Duration / Buffer','Calendar Group']
for i,h in enumerate(hdrs):
    c = cal_t.rows[0].cells[i]; shade_cell(c,'1E293B')
    pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    pp.paragraph_format.space_before=Pt(8); pp.paragraph_format.space_after=Pt(8)
    rr=pp.add_run(h); rr.bold=True; rr.font.size=Pt(9); rr.font.color.rgb=WHITE

cal_rows = [
    ('CAL-01','Audio Recording','Non-Member','2 hrs / 30 min','Non-Member Group','F8FAFC','475569'),
    ('CAL-02','Audio Recording','Gold',      '2 hrs / 30 min','Gold Group',      'FFFBEB','D97706'),
    ('CAL-03','Audio Recording','VIP',       '2 hrs / 30 min','VIP Group',       'F5F3FF','7C3AED'),
    ('CAL-04','Video Podcast',  'Non-Member','2 hrs / 30 min','Non-Member Group','F8FAFC','475569'),
    ('CAL-05','Video Podcast',  'Gold',      '2 hrs / 30 min','Gold Group',      'FFFBEB','D97706'),
    ('CAL-06','Video Podcast',  'VIP',       '2 hrs / 30 min','VIP Group',       'F5F3FF','7C3AED'),
    ('CAL-07','Live Streaming', 'Non-Member','2 hrs / 30 min','Non-Member Group','F8FAFC','475569'),
    ('CAL-08','Live Streaming', 'Gold',      '2 hrs / 30 min','Gold Group',      'FFFBEB','D97706'),
    ('CAL-09','Live Streaming', 'VIP',       '2 hrs / 30 min','VIP Group',       'F5F3FF','7C3AED'),
]
for ri,(cid,svc,tier,dur,grp,bg,tc) in enumerate(cal_rows):
    row = cal_t.rows[ri+1]
    for ci,val in enumerate([cid,svc,tier,dur,grp]):
        c = row.cells[ci]
        shade_cell(c, bg if ci==2 else ('FFFFFF' if ri%2==0 else 'F8FAFC'))
        pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_before=Pt(7); pp.paragraph_format.space_after=Pt(7)
        rr=pp.add_run(val); rr.font.size=Pt(10)
        if   ci==0: rr.bold=True; rr.font.color.rgb=BLUE
        elif ci==2: rr.bold=True; rr.font.color.rgb=rgb(tc)
        else:                      rr.font.color.rgb=GRAY


# ═══════════════════════════════════════════════
#  SECTION 4 — AVAILABILITY
# ═══════════════════════════════════════════════
gap(doc, 16)
section_hdr(doc, '🔒  HOW THE SYSTEM PREVENTS DOUBLE-BOOKINGS')
p(doc, 'Three independent layers work together. All three must be available simultaneously for a time slot to appear.',
  size=10.5, color=GRAY, sb=0, sa=14)

avail = [
    ('01','👥  3 Staff Slots — Concurrent Capacity',[
        'Three placeholder "studio slots" represent the maximum number of simultaneous bookings',
        'Since there are 3 studios, maximum 3 sessions can run at the same time — across ALL tiers combined',
        'When all 3 slots are occupied, no further bookings can be made for that time period',
        'Example: 3 bookings exist at 2pm. A 4th person (any tier) tries to book 2pm → all slots claimed → slot hidden',
    ]),
    ('02','🏠  3 Rooms — Studio A, B & C',[
        'Each physical studio is registered as a Room in GoHighLevel',
        'Each room holds exactly 1 booking at a time',
        'When a booking is confirmed, GHL auto-assigns an available room and claims it instantly',
        'The room is blocked across ALL calendars and ALL tiers simultaneously',
        'Example: VIP books Studio A at 3pm → Studio A blocked for Non-Member and Gold at 3pm too',
    ]),
    ('03','🎛️  2 Equipment Types — Shared Gear',[
        'Teleprompter (qty: 2) — linked to Video Podcast and Live Streaming calendars only',
        'Streaming Encoder (qty: 1) — linked to Live Streaming calendars only',
        'Running out of equipment blocks that service even if staff slots and rooms are free',
        'Example: 2 Video Podcast sessions at 4pm → both teleprompters claimed → 3rd Video Podcast at 4pm BLOCKED',
    ]),
]
for num,title,bullets in avail:
    pp = doc.add_paragraph(); pp.paragraph_format.space_before = Pt(10)
    rn=pp.add_run(f'{num}  '); rn.bold=True; rn.font.size=Pt(18); rn.font.color.rgb=BLUE
    rt=pp.add_run(title);      rt.bold=True; rt.font.size=Pt(11.5); rt.font.color.rgb=DARK
    rv=pp.add_run('  ✅');     rv.font.size=Pt(11)
    for b in bullets:
        pb=doc.add_paragraph(style='List Bullet')
        pb.paragraph_format.left_indent=Cm(1.5); pb.paragraph_format.space_after=Pt(3)
        rb=pb.add_run(b); rb.font.size=Pt(10); rb.font.color.rgb=GRAY

gap(doc, 10)
p(doc, 'Live Availability Example — Tuesday 2:00 PM', size=11, bold=True, color=DARK, sb=8, sa=0)

ex_lines = [
    ('Booking 1  —  Non-Member  |  Audio Recording  |  2pm', DARK,  True),
    ('   ✓  Staff Slot 1 free  →  assigned',                  GREEN, False),
    ('   ✓  Studio A free  →  assigned & locked',             GREEN, False),
    ('   ✓  CONFIRMED',                                        GREEN, True),
    ('',None,False),
    ('Booking 2  —  Gold Member  |  Video Podcast  |  2pm',  DARK,  True),
    ('   ✓  Staff Slot 2 free  →  assigned',                  GREEN, False),
    ('   ✓  Studio B free  →  assigned & locked',             GREEN, False),
    ('   ✓  Teleprompter #1 free  →  assigned & locked',      GREEN, False),
    ('   ✓  CONFIRMED',                                        GREEN, True),
    ('',None,False),
    ('Booking 3  —  VIP Member  |  Live Streaming  |  2pm',  DARK,  True),
    ('   ✓  Staff Slot 3 free  →  assigned',                  GREEN, False),
    ('   ✓  Studio C free  →  assigned & locked',             GREEN, False),
    ('   ✓  Streaming Encoder free  →  assigned & locked',    GREEN, False),
    ('   ✓  CONFIRMED',                                        GREEN, True),
    ('',None,False),
    ('Booking 4 (attempted)  —  Any Tier  |  Any Service  |  2pm', DARK, True),
    ('   ✗  All 3 staff slots occupied',                       RED,  False),
    ('   ✗  All 3 studios occupied',                           RED,  False),
    ('   ✗  SLOT NOT SHOWN  —  customer sees next available time', RED, True),
    ('',None,False),
    ('→  Next available:  4:30 PM  (2hr session + 30min buffer)', LGRAY, False),
]
for line,color,bold in ex_lines:
    ep=doc.add_paragraph()
    ep.paragraph_format.space_before=Pt(0); ep.paragraph_format.space_after=Pt(1)
    ep.paragraph_format.left_indent=Cm(0.4); shade_para(ep,'F8FAFC')
    if line:
        er=ep.add_run(line); er.font.size=Pt(10); er.font.name='Courier New'; er.bold=bold
        if color: er.font.color.rgb=color


# ═══════════════════════════════════════════════
#  SECTION 5 — BOOKING LINKS (clickable, highlighted)
# ═══════════════════════════════════════════════
gap(doc, 16)
section_hdr(doc, '🔗  LIVE BOOKING PAGES — TIER-SPECIFIC LINKS')
p(doc, 'Each tier has its own dedicated booking page. The member portal routes each member to their correct page '
       'automatically based on their GHL contact tag. No member can see another tier\'s pricing.',
  size=10.5, color=GRAY, sb=0, sa=14)

booking_pages = [
    ('01','NON-MEMBER BOOKING PAGE','Tier 1 — Standard Rates  ·  Tag: member-non',
     'https://api.leadconnectorhq.com/widget/service-menu/6a2d5353b5de14a8a631a28c',
     'F1F5F9','475569','CBD5E1'),
    ('02','GOLD MEMBER BOOKING PAGE','Tier 2 — Member Rates  ·  Tag: member-gold',
     'https://api.leadconnectorhq.com/widget/service-menu/6a2d5371e1adc836daa87e36',
     'FEF3C7','D97706','FCD34D'),
    ('03','VIP MEMBER BOOKING PAGE','Tier 3 — Premium Rates  ·  Tag: member-vip',
     'https://api.leadconnectorhq.com/widget/service-menu/6a2d538fbd4b5f033b642719',
     'EDE9FE','7C3AED','C4B5FD'),
]

for num,title,subtitle,url,bg,tc,border in booking_pages:
    bt = doc.add_table(rows=1, cols=1); no_borders(bt)
    c = bt.rows[0].cells[0]; shade_cell(c, bg)

    # Add left border accent via XML
    tcPr = c._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left_b = OxmlElement('w:left')
    left_b.set(qn('w:val'), 'single'); left_b.set(qn('w:sz'), '24')
    left_b.set(qn('w:color'), tc)
    tcBorders.append(left_b); tcPr.append(tcBorders)

    pp1 = c.paragraphs[0]; pp1.paragraph_format.space_before = Pt(14)
    r1a = pp1.add_run(f'{num}  ')
    r1a.bold=True; r1a.font.size=Pt(16); r1a.font.color.rgb=rgb(tc)
    r1b = pp1.add_run(title)
    r1b.bold=True; r1b.font.size=Pt(13); r1b.font.color.rgb=rgb(tc)

    pp2 = c.add_paragraph()
    r2=pp2.add_run(subtitle); r2.font.size=Pt(10); r2.font.color.rgb=GRAY

    pp3 = c.add_paragraph()
    r3=pp3.add_run('Services: Audio Recording  ·  Video Podcast  ·  Live Streaming')
    r3.font.size=Pt(10); r3.font.color.rgb=GRAY

    # Clickable link with instruction
    pp4 = c.add_paragraph(); pp4.paragraph_format.space_after = Pt(4)
    r4a = pp4.add_run('👉  Click here to explore the service menu:  ')
    r4a.bold=True; r4a.font.size=Pt(10); r4a.font.color.rgb=DARK
    add_hyperlink(pp4, url, url, color_hex=tc, bold=False, size=10)

    gap(doc, 8)


# ═══════════════════════════════════════════════
#  SECTION 6 — CHECKLIST
# ═══════════════════════════════════════════════
gap(doc, 16)
section_hdr(doc, '✅  WHAT WAS BUILT — PHASE 1 CHECKLIST')

checklist = [
    ('3 Staff Slots',       'Studio Slot 1, 2, 3 — each representing one concurrent booking capacity — availability set to studio opening hours'),
    ('9 Service Calendars', 'CAL-01 through CAL-09 — all 3 service types across all 3 tiers — 2hr duration, 30min buffer each'),
    ('3 Rooms (Studios)',   'Studio A, Studio B, Studio C — linked to all 9 calendars — auto-assigned on booking — cross-tier double-booking prevented'),
    ('2 Equipment Items',   'Teleprompter (qty 2, linked to Video + Streaming) and Streaming Encoder (qty 1, linked to Live Streaming only)'),
    ('3 Calendar Groups',   'Non-Member Group, Gold Group, VIP Group — each containing exactly 3 calendars for that tier only'),
    ('3 Service Menus',     'Non-Member, Gold, and VIP booking pages — all 3 live, tested, and linked above'),
]
for title,desc in checklist:
    cp=doc.add_paragraph()
    cp.paragraph_format.space_before=Pt(7); cp.paragraph_format.space_after=Pt(4)
    cr1=cp.add_run('✓  '); cr1.bold=True; cr1.font.size=Pt(13); cr1.font.color.rgb=GREEN
    cr2=cp.add_run(title+' — '); cr2.bold=True; cr2.font.size=Pt(11); cr2.font.color.rgb=DARK
    cr3=cp.add_run(desc); cr3.font.size=Pt(10.5); cr3.font.color.rgb=GRAY


# ═══════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════
out = 'd:/Watts Company/GHL_SNAPSHOT_BRAIN/Imp_Documents/Podcast_Studio_Booking_System_v2.docx'
doc.save(out)
print(f'Saved: {out}')
